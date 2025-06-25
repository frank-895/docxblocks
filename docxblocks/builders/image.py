from docx.shared import Inches
from PIL import Image as PILImage
import os

class ImageBuilder:
    @staticmethod
    def build(doc, image_path=None, parent=None, index=None, **style_kwargs):
        if not image_path or not os.path.isfile(image_path):
            para = doc.add_paragraph("Image not found")
            parent.insert(index, para._element)
            return

        new_para = doc.add_paragraph()
        run = new_para.add_run()

        try:
            with PILImage.open(image_path) as img:
                width_px, height_px = img.size
                dpi_x, dpi_y = img.info.get("dpi", (72, 72))
                width_in = width_px / dpi_x
                height_in = height_px / dpi_y

                scale = 1.0
                max_width = _parse_measurement(style_kwargs.get("max_width"))
                max_height = _parse_measurement(style_kwargs.get("max_height"))

                if max_width:
                    scale = min(scale, max_width / width_in)
                if max_height:
                    scale = min(scale, max_height / height_in)

                run.add_picture(image_path, width=Inches(width_in * scale), height=Inches(height_in * scale))

        except Exception as e:
            error_para = doc.add_paragraph(f"[Image failed to render: {e}]")
            parent.insert(index, error_para._element)
            return

        parent.insert(index, new_para._element)


def _parse_measurement(value):
    """
    Accepts strings like '4in' or '300px' and returns inches as float.
    Supports only inches and pixels for now.
    """
    if not value or not isinstance(value, str):
        return None

    value = value.strip().lower()

    if value.endswith("in"):
        try:
            return float(value.replace("in", ""))
        except ValueError:
            return None
    elif value.endswith("px"):
        try:
            px = float(value.replace("px", ""))
            return px / 96.0  # assuming 96 dpi standard
        except ValueError:
            return None
    else:
        return None
