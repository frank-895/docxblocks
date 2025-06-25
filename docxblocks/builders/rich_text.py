from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docxblocks.schema.blocks import Block, TextBlock, HeadingBlock, BulletBlock, TableBlock, ImageBlock
from docxblocks.utils.styles import apply_style_to_run, set_paragraph_alignment
from docxblocks.builders.table import TableBuilder
from docxblocks.builders.image import ImageBuilder


class RichTextBuilder:
    def __init__(self, doc, parent, index):
        self.doc = doc
        self.parent = parent
        self.index = index

    def render(self, blocks: list):
        """Render a list of validated or raw block dicts into the document."""
        validated_blocks = [Block.parse_obj(b) for b in blocks]
        for block in validated_blocks:
            if isinstance(block, TextBlock):
                self._render_text(block)
            elif isinstance(block, HeadingBlock):
                self._render_heading(block)
            elif isinstance(block, BulletBlock):
                self._render_bullets(block)
            elif isinstance(block, TableBlock):
                self._render_table(block)
            elif isinstance(block, ImageBlock):
                self._render_image(block)

    def _render_text(self, block: TextBlock):
        lines = block.text.split("\n")
        for line in lines:
            para = self.doc.add_paragraph(
                style=block.style.style if block.style and block.style.style else "Normal"
            )
            run = para.add_run(line)
            apply_style_to_run(run, block.style)
            set_paragraph_alignment(para, block.style.align if block.style else None)
            self.parent.insert(self.index, para._element)
            self.index += 1

    def _render_heading(self, block: HeadingBlock):
        style_name = (
            block.style.style if block.style and block.style.style else f"Heading {block.level}"
        )
        para = self.doc.add_paragraph(style=style_name)
        run = para.add_run(block.text)
        apply_style_to_run(run, block.style)
        set_paragraph_alignment(para, block.style.align if block.style else None)
        self.parent.insert(self.index, para._element)
        self.index += 1

    def _render_bullets(self, block: BulletBlock):
        for item in block.items:
            para = self.doc.add_paragraph(style="Report Bullet")
            run = para.add_run(item)
            apply_style_to_run(run, block.style)
            set_paragraph_alignment(para, block.style.align if block.style else None)
            self.parent.insert(self.index, para._element)
            self.index += 1

    def _render_table(self, block: TableBlock):
        TableBuilder.build(
            self.doc,
            placeholder=None,
            content=block.content,
            parent=self.parent,
            index=self.index,
            **(block.style.dict() if block.style else {})
        )
        self.index += 1

    def _render_image(self, block: ImageBlock):
        ImageBuilder.build(
            self.doc,
            placeholder=None,
            image_path=block.path,
            parent=self.parent,
            index=self.index,
            **(block.style.dict() if block.style else {})
        )
        self.index += 1