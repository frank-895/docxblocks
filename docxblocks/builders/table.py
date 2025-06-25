from docx.shared import RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.text import WD_ALIGN_PARAGRAPH

class TableBuilder:
    @staticmethod
    def build(doc, placeholder=None, content=None, parent=None, index=None, **style_kwargs):
        if not parent or index is None or not content:
            return

        headers = content.get("headers", [])
        rows = content.get("rows", [])

        num_cols = len(headers) if headers else len(rows[0])
        table = doc.add_table(rows=0, cols=num_cols)
        table.style = "Table Grid"

        column_widths = style_kwargs.get("column_widths")

        if column_widths:
            total_width = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin
            for i, fraction in enumerate(column_widths):
                if i < len(table.columns):
                    table.columns[i].width = int(total_width * fraction)

        if headers:
            row = table.add_row().cells
            for i, header_text in enumerate(headers):
                cell = row[i]
                para = cell.paragraphs[0]
                run = para.add_run(str(header_text))
                run.font.bold = True

                header_styles = style_kwargs.get("header_styles", {})
                _apply_cell_style(cell, para, run, header_styles)

        for row_data in rows:
            row = table.add_row().cells
            for i, cell_val in enumerate(row_data):
                cell = row[i]
                para = cell.paragraphs[0]
                run = para.add_run(str(cell_val))
                run.font.bold = False

                col_styles = style_kwargs.get("column_styles", {}).get(i, {})
                _apply_cell_style(cell, para, run, col_styles)

                cell_styles = style_kwargs.get("cell_styles", {}).get((len(table.rows)-2, i), {})
                _apply_cell_style(cell, para, run, cell_styles)

        parent.insert(index, table._element)


def _apply_cell_style(cell, para, run, styles):
    if styles.get("align"):
        _set_paragraph_alignment(para, styles["align"])
    if styles.get("bold"):
        run.font.bold = True
    if styles.get("bg_color"):
        _set_cell_bg_color(cell, styles["bg_color"])
    if styles.get("font_color"):
        run.font.color.rgb = RGBColor.from_string(styles["font_color"])


def _set_cell_bg_color(cell, hex_color):
    cell_xml = cell._tc
    props = cell_xml.get_or_add_tcPr()
    props.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'))


def _set_paragraph_alignment(paragraph, align):
    alignment_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }
    if align in alignment_map:
        paragraph.alignment = alignment_map[align]
