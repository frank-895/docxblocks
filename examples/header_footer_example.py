from docx import Document
from docxblocks.core.inserter import DocxBuilder
from PIL import Image
import tempfile
import os

# Create a simple template
doc = Document()
doc.add_paragraph("{{main}}")
doc.save("header_footer_template.docx")

# Define blocks including headers and footers
blocks = [
    # Document header that appears on all pages
    {
        "type": "header",
        "apply_to": "all",
        "content": [
            {
                "type": "text",
                "text": "Company Name\tDocument Title\tPage {{page}}",
                "style": {"align": "center", "font_color": "666666"}
            }
        ]
    },
    
    # Document footer that appears on all pages
    {
        "type": "footer",
        "apply_to": "all", 
        "content": [
            {
                "type": "text",
                "text": "© 2024 Company Name. All rights reserved.\tConfidential\t{{date}}",
                "style": {"align": "center", "font_color": "999999", "italic": True}
            }
        ]
    },
    
    # Main content
    {
        "type": "heading",
        "text": "Report with Headers and Footers",
        "level": 1,
        "style": {"align": "center", "bold": True}
    },
    {
        "type": "text",
        "text": "This document demonstrates the header and footer functionality. The header contains company information and page numbers, while the footer includes copyright and date information."
    },
    {
        "type": "page_break"
    },
    {
        "type": "heading",
        "text": "Second Page",
        "level": 2
    },
    {
        "type": "text",
        "text": "This is the second page of the document. Notice that the header and footer appear on all pages as configured."
    }
]

# Build the document
builder = DocxBuilder("header_footer_template.docx")
builder.insert("{{main}}", blocks)
builder.save("header_footer_output.docx")
print("Header/Footer example generated: header_footer_output.docx")


# Example 2: Different headers for first, odd, and even pages
print("\nGenerating advanced header/footer example...")

# Create another template
doc2 = Document()
doc2.add_paragraph("{{content}}")
doc2.save("advanced_header_footer_template.docx")

advanced_blocks = [
    # Special header for first page only
    {
        "type": "header", 
        "apply_to": "first",
        "content": [
            {
                "type": "text",
                "text": "FIRST PAGE HEADER - Report Title",
                "style": {"align": "center", "bold": True, "font_color": "000080"}
            }
        ]
    },
    
    # Header for odd pages (after first page)
    {
        "type": "header",
        "apply_to": "odd", 
        "content": [
            {
                "type": "text",
                "text": "Odd Page Header\t\tPage {{page}}",
                "style": {"align": "left", "font_color": "333333"}
            }
        ]
    },
    
    # Header for even pages
    {
        "type": "header",
        "apply_to": "even",
        "content": [
            {
                "type": "text", 
                "text": "Page {{page}}\t\tEven Page Header",
                "style": {"align": "right", "font_color": "333333"}
            }
        ]
    },
    
    # Different footer for first page
    {
        "type": "footer",
        "apply_to": "first",
        "content": [
            {
                "type": "text",
                "text": "Report generated on {{date}} - Page 1 of {{total_pages}}",
                "style": {"align": "center", "italic": True}
            }
        ]
    },
    
    # Footer for all other pages
    {
        "type": "footer", 
        "apply_to": "odd",
        "content": [
            {
                "type": "text",
                "text": "© Company\tConfidential\tPage {{page}}",
                "style": {"align": "center", "font_color": "666666"}
            }
        ]
    },
    
    {
        "type": "footer",
        "apply_to": "even", 
        "content": [
            {
                "type": "text",
                "text": "Page {{page}}\tConfidential\t© Company",
                "style": {"align": "center", "font_color": "666666"}
            }
        ]
    },
    
    # Main content
    {
        "type": "heading",
        "text": "Advanced Header/Footer Demo", 
        "level": 1,
        "style": {"align": "center"}
    },
    {
        "type": "text",
        "text": "This document demonstrates different headers and footers for first page, odd pages, and even pages. Notice how each page type has its unique header and footer design."
    },
    {
        "type": "page_break"
    },
    {
        "type": "heading",
        "text": "Page 2 (Even)",
        "level": 2
    },
    {
        "type": "text", 
        "text": "This is an even-numbered page. It has different header and footer styling."
    },
    {
        "type": "page_break"
    },
    {
        "type": "heading",
        "text": "Page 3 (Odd)",
        "level": 2
    },
    {
        "type": "text",
        "text": "This is an odd-numbered page. It uses the odd page header and footer styles."
    }
]

# Build the advanced document
builder2 = DocxBuilder("advanced_header_footer_template.docx")
builder2.insert("{{content}}", advanced_blocks)
builder2.save("advanced_header_footer_output.docx")
print("Advanced Header/Footer example generated: advanced_header_footer_output.docx")


# Example 3: Headers and footers with images
print("\nGenerating header/footer with images example...")

# Create a test image for the example
with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_file:
    img = Image.new('RGB', (150, 50), color='blue')
    img.save(tmp_file.name, dpi=(96, 96))
    logo_path = tmp_file.name

# Create another template
doc3 = Document()
doc3.add_paragraph("{{content}}")
doc3.save("header_footer_images_template.docx")

image_blocks = [
    # Header with image and text
    {
        "type": "header",
        "apply_to": "all",
        "content": [
            {
                "type": "image",
                "path": logo_path,
                "style": {"max_width": "1in"}
            },
            {
                "type": "text",
                "text": "Company Report",
                "style": {"align": "center", "bold": True, "font_color": "000080"}
            }
        ]
    },
    
    # Footer with image
    {
        "type": "footer",
        "apply_to": "all",
        "content": [
            {
                "type": "text",
                "text": "© 2024 Company",
                "style": {"align": "left", "font_color": "666666"}
            },
            {
                "type": "image",
                "path": logo_path,
                "style": {"max_width": "0.5in"}
            },
            {
                "type": "text",
                "text": "Page {{page}}",
                "style": {"align": "right", "font_color": "666666"}
            }
        ]
    },
    
    # Main content
    {
        "type": "heading",
        "text": "Header/Footer with Images Demo",
        "level": 1,
        "style": {"align": "center"}
    },
    {
        "type": "text",
        "text": "This document demonstrates headers and footers that contain images. The header includes a company logo and title, while the footer has a smaller logo with page numbers."
    },
    {
        "type": "page_break"
    },
    {
        "type": "heading",
        "text": "Second Page",
        "level": 2
    },
    {
        "type": "text",
        "text": "Notice that the images appear correctly in both the header and footer on all pages."
    }
]

# Build the document with images
builder3 = DocxBuilder("header_footer_images_template.docx")
builder3.insert("{{content}}", image_blocks)
builder3.save("header_footer_images_output.docx")

# Clean up the temporary image
os.unlink(logo_path)

print("Header/Footer with Images example generated: header_footer_images_output.docx")
print("Note: Images in headers and footers are now fully supported!") 